# banksoal_v3.py
import os
import pandas as pd
import random
from fpdf import FPDF
from docx import Document

BANK_FILE = "database.csv"

class QuestionBank:
    def __init__(self):
        self.questions = pd.DataFrame()
        self.original_questions = pd.DataFrame()  # نسخه اصلی بانک
        self.selected_questions = pd.DataFrame()
        self.load_questions()

    def load_questions(self):
        try:
            if os.path.exists(BANK_FILE):
                self.questions = pd.read_csv(
                    BANK_FILE, dtype=str, encoding="utf-8-sig", on_bad_lines="skip"
                ).fillna("")
            else:
                self.questions = pd.DataFrame(columns=[
                    "Question", "Answer", "Option1", "Option2", "Option3", "Category", "Source"
                ])
            self.original_questions = self.questions.copy()
        except Exception:
            self.questions = pd.DataFrame(columns=[
                "Question", "Answer", "Option1", "Option2", "Option3", "Category", "Source"
            ])
            self.original_questions = self.questions.copy()

    def save_questions(self):
        try:
            self.questions.to_csv(BANK_FILE, index=False, encoding="utf-8-sig")
            self.original_questions = self.questions.copy()
            return True, ""
        except Exception as e:
            return False, str(e)

    def add_question(self, question_data):
        try:
            row = pd.DataFrame([question_data], columns=self.questions.columns)
            self.questions = pd.concat([self.questions, row], ignore_index=True)
            self.original_questions = self.questions.copy()
            return True, ""
        except Exception as e:
            return False, str(e)

    def list_questions(self):
        return self.questions.copy()

    def filter_questions(self, **kwargs):
        filtered = self.questions.copy()
        for key, value in kwargs.items():
            if key in filtered.columns and value:
                filtered = filtered[filtered[key].astype(str).str.contains(str(value), case=False, na=False)]
        return filtered

    def select_questions(self, mode="manual", indices=None, count=None, indices_are_labels=False):
        try:
            if mode == "manual" and indices:
                if indices_are_labels:
                    valid_labels = [i for i in indices if i in self.questions.index]
                    self.selected_questions = self.questions.loc[valid_labels].copy()
                else:
                    max_n = len(self.questions)
                    safe = [i for i in indices if 0 <= i < max_n]
                    self.selected_questions = (
                        self.questions.iloc[safe].copy() if safe else pd.DataFrame()
                    )
            elif mode == "random" and count:
                cnt = min(int(count), max(0, len(self.questions)))
                self.selected_questions = (
                    self.questions.sample(n=cnt).copy() if cnt > 0 else pd.DataFrame()
                )
            else:
                self.selected_questions = pd.DataFrame()
            return True, ""
        except Exception as e:
            self.selected_questions = pd.DataFrame()
            return False, str(e)

    def export_pdf(self, filename="output.pdf"):
        try:
            data = self.selected_questions if not self.selected_questions.empty else self.questions
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Arial", size=12)

            for idx, row in data.reset_index(drop=True).iterrows():
                qnum = idx + 1
                pdf.multi_cell(0, 8, f"{qnum}. {row.get('Question', '')}")
                o1, o2, o3 = row.get('Option1', ''), row.get('Option2', ''), row.get('Option3', '')
                if o1 or o2 or o3:
                    pdf.multi_cell(0, 8, f"A) {o1}    B) {o2}    C) {o3}")
                pdf.ln(4)
            pdf.output(filename)
            return True, ""
        except Exception as e:
            return False, str(e)

    def export_word(self, filename="output.docx"):
        try:
            data = self.selected_questions if not self.selected_questions.empty else self.questions
            doc = Document()
            for idx, row in data.reset_index(drop=True).iterrows():
                qnum = idx + 1
                doc.add_paragraph(f"{qnum}. {row.get('Question', '')}")
                o1, o2, o3 = row.get('Option1', ''), row.get('Option2', ''), row.get('Option3', '')
                if o1 or o2 or o3:
                    doc.add_paragraph(f"A) {o1}    B) {o2}    C) {o3}")
                doc.add_paragraph("")
            doc.save(filename)
            return True, ""
        except Exception as e:
            return False, str(e)

    def reset_view(self):
        self.questions = self.original_questions.copy()
        self.selected_questions = pd.DataFrame()

    def label_to_position(self, label):
        try:
            labels = list(self.questions.index)
            return labels.index(label)
        except ValueError:
            return None
